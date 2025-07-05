
import streamlit as st
import tempfile
import os
import fitz  # PyMuPDF
import pandas as pd
import torch
import re
import json
import pypandoc
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

device = "cuda" if torch.cuda.is_available() else "cpu"

@st.cache_resource
def load_summarizer(model_name):
    return pipeline("summarization", model=model_name, device=0 if torch.cuda.is_available() else -1)

@st.cache_resource
def load_rewriter(model_name):
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name).to(device)
    return pipeline("text2text-generation", model=model, tokenizer=tokenizer, device=0 if torch.cuda.is_available() else -1)

HEADERS = [
    # Standard sections
    "Abstract", "Introduction", "Background",
    "Related Work", "Literature Review", "Theoretical Framework",
    "Methodology", "Research Methodology", "Materials and Methods",
    "Experimental Setup", "Experimental Design", "Research Design",
    "Implementation", "System Architecture", "Algorithm",
    "Data Collection", "Data Analysis", "Statistical Analysis",
    "Results", "Findings", "Experimental Results", "Case Study",
    "Discussion", "Analysis", "Interpretation", "Evaluation",
    "Conclusion", "Conclusions", "Summary", "Future Work",
    "Recommendations", "Implications", "Limitations",
    "References", "Bibliography", "Acknowledgements",

    # Subject-specific sections
    # STEM
    "Hypothesis", "Procedure", "Observations",
    "Calculations", "Derivation", "Proof",
    "Simulation", "Modeling", "Validation",
    "Performance Metrics", "Benchmarking",

    # Medical/Health Sciences
    "Clinical Trial", "Patient Characteristics",
    "Intervention", "Outcome Measures",
    "Adverse Effects", "Ethical Considerations",

    # Social Sciences
    "Research Questions", "Conceptual Framework",
    "Participant Demographics", "Survey Instrument",
    "Qualitative Analysis", "Quantitative Analysis",
    "Themes", "Coding Scheme",

    # Business/Economics
    "Market Analysis", "Financial Projections",
    "SWOT Analysis", "Case Analysis",
    "Return on Investment", "Cost-Benefit Analysis",

    # Humanities
    "Historical Context", "Textual Analysis",
    "Critical Review", "Comparative Analysis",
    "Theoretical Perspective",

    # Engineering/CS
    "System Requirements", "Technical Specifications",
    "Pseudocode", "Flowchart", "UML Diagram",
    "User Interface", "API Documentation",
    "Error Analysis", "Optimization",

    # Appendices
    "Appendix A", "Appendix B", "Supplementary Materials",
    "Additional Results", "Raw Data",

    # Conference/Journal Specific
    "Key Takeaways", "Author Contributions",
    "Conflict of Interest", "Data Availability",
    "Supplementary Information"
]
header_pattern = re.compile(rf"\n*(?:\d{{0,2}}[\.\)]?\s*)?({'|'.join(HEADERS)})\s*\n", re.IGNORECASE)

def extract_layout_text(filepath):
    doc = fitz.open(filepath)
    full_text = ""
    for page in doc:
        blocks = page.get_text("blocks")
        sorted_blocks = sorted(blocks, key=lambda b: (round(b[1], 1), round(b[0], 1)))
        for b in sorted_blocks:
            text = b[4].strip()
            if text:
                full_text += text + "\n"
    return full_text

def extract_sections(text):
    matches = list(header_pattern.finditer(text))
    sections = {}
    for i, match in enumerate(matches):
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        header = match.group(1).strip().title()
        content = text[start:end].strip()
        sections[header] = content
    return sections

def summarize_text(text, summarizer, max_length=130, min_length=30):
    try:
        if len(text.split()) > 1024:
            chunks = [text[i:i+1024] for i in range(0, len(text), 1024)]
            summaries = []
            for chunk in chunks:
                summary = summarizer(chunk, max_length=max_length, min_length=min_length, do_sample=False)
                summaries.append(summary[0]['summary_text'])
            return " ".join(summaries)
        else:
            result = summarizer(text, max_length=max_length, min_length=min_length, do_sample=False)
            return result[0]['summary_text']
    except Exception as e:
        return text[:max_length]

def rewrite_text(text, rewriter, temperature=0.7, chunk_size=400):
    try:
        if len(text) > chunk_size:
            chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
            rewritten_chunks = []
            for chunk in chunks:
                result = rewriter(f"paraphrase: {chunk}", do_sample=True, temperature=temperature, num_return_sequences=1)
                rewritten_chunks.append(result[0]['generated_text'])
            return " ".join(rewritten_chunks)
        else:
            result = rewriter(f"paraphrase: {text}", do_sample=True, temperature=temperature, num_return_sequences=1)
            return result[0]['generated_text']
    except:
        return text

def convert_docx_to_pdf(docx_path):
    output_path = docx_path.replace(".docx", ".pdf")
    try:
        pypandoc.convert_file(docx_path, 'pdf', outputfile=output_path)
        return output_path
    except:
        raise RuntimeError("DOCX to PDF conversion failed. Ensure LaTeX is available.")

def parse_to_df(filepath):
    raw_text = extract_layout_text(filepath)
    sections = extract_sections("\n" + raw_text + "\n")
    return pd.DataFrame(sections.items(), columns=["Section", "Content"])

def create_docx(text, filename):
    doc = Document()
    for section in text.split("### "):
        if section.strip():
            lines = section.strip().split("\n", 1)
            doc.add_heading(lines[0].strip(), level=1)
            doc.add_paragraph(lines[1].strip() if len(lines) > 1 else "")
    doc.save(filename)

def create_pdf(text, filename):
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    for section in text.strip().split("### "):
        if section.strip():
            parts = section.strip().split("\n", 1)
            header = parts[0].strip()
            body = parts[1].strip() if len(parts) > 1 else ""
            story.append(Paragraph(f"<b>{header}</b>", styles["Heading2"]))
            story.append(Spacer(1, 6))
            story.append(Paragraph(body.replace("\n", "<br/>"), styles["BodyText"]))
            story.append(Spacer(1, 12))

    doc.build(story)

# ---------------- UI ----------------
st.title("📄 Research Paper Summarizer and Rewriter")

uploaded_file = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])
if uploaded_file:
    st.session_state["file"] = uploaded_file

summarizer_options = {
    "facebook/bart-large-cnn (default)": "facebook/bart-large-cnn",
    "sshleifer/distilbart-cnn-12-6 (fast)": "sshleifer/distilbart-cnn-12-6",
    "google/pegasus-cnn_dailymail (high-quality)": "google/pegasus-cnn_dailymail",
    "Falconsai/text_summarization (lightweight)": "Falconsai/text_summarization"
}
selected_model_name = st.selectbox("Choose summarizer model:", list(summarizer_options.keys()))
selected_model = summarizer_options[selected_model_name]

rewriter_options = {
    "Vamsi/T5_Paraphrase_Paws (fastest)": "Vamsi/T5_Paraphrase_Paws",
    "ramsrigouthamg/t5_paraphraser": "ramsrigouthamg/t5_paraphraser",
    "Vamsi/pegasus_paraphrase": "Vamsi/pegasus_paraphrase",
    "prithivida/parrot_paraphraser_on_T5 (default)": "prithivida/parrot_paraphraser_on_T5"
}
selected_rewriter_name = st.selectbox("Choose rewriter model:", list(rewriter_options.keys()))
selected_rewriter = rewriter_options[selected_rewriter_name]

apply_rewriting = st.checkbox("Apply rewriting before summarizing", value=True)

if uploaded_file and selected_model and selected_rewriter:
    if st.button("🟢 Process"):
        st.session_state["selected_model"] = selected_model
        st.session_state["selected_rewriter"] = selected_rewriter
        st.session_state["apply_rewriting"] = apply_rewriting
        st.session_state["trigger_process"] = True

        # 🔥 CLEAR previous final output so processing block runs again
        if "final_output" in st.session_state:
            del st.session_state["final_output"]

        st.rerun()

if (
    "file" in st.session_state and
    "trigger_process" in st.session_state and
    st.session_state["trigger_process"] and
    "final_output" not in st.session_state
):
    uploaded_file = st.session_state["file"]
    selected_model = st.session_state["selected_model"]
    selected_rewriter = st.session_state["selected_rewriter"]
    apply_rewriting = st.session_state["apply_rewriting"]

    summarizer = load_summarizer(selected_model)
    rewriter = load_rewriter(selected_rewriter)

    with st.spinner("Processing file..."):
        temp_input_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
        with open(temp_input_path, "wb") as f:
            f.write(uploaded_file.read())

        filepath = temp_input_path
        if uploaded_file.name.endswith(".docx"):
            try:
                filepath = convert_docx_to_pdf(temp_input_path)
            except Exception as e:
                st.error(f"Failed to convert DOCX to PDF: {e}")
                st.stop()

        try:
            st.write("✅ Parsing file...")
            df = parse_to_df(filepath)

            if apply_rewriting:
                st.write("✅ Rewriting content...")
                df['rewritten_content'] = df['Content'].apply(lambda x: rewrite_text(x, rewriter, temperature=0.5))
                content_for_summary = df['rewritten_content']
            else:
                st.write("⚠️ Skipping rewriting step...")
                content_for_summary = df['Content']

            st.write("✅ Summarizing content...")
            df['summerized_content'] = content_for_summary.apply(lambda x: summarize_text(x, summarizer))
            st.write("✅ Done.")

            output_text = ""
            for _, row in df.iterrows():
                output_text += f"### {row['Section']}\n\n"
                # Split summary into sentences and bullet them
                bullets = row['summerized_content'].split(". ")
                for bullet in bullets:
                    bullet = bullet.strip().rstrip(".")
                    if bullet:
                        output_text += f"- {bullet}.\n"
                output_text += "\n"

            st.session_state["final_output"] = output_text

        except Exception as e:
            st.error(f"Error processing file: {e}")
            st.stop()

# ---------------- Editable Output & Download ----------------
if "final_output" in st.session_state:
    st.subheader("📝 Final Output (Editable)")

    with st.form("edit_and_download"):
        edited_output = st.text_area("Edit the summarized content below:", st.session_state["final_output"], height=600)
        submit = st.form_submit_button("💾 Save and Show Download Buttons")

    if submit:
        st.session_state["final_output"] = edited_output

        st.subheader("⬇️ Download Your File")

        docx_path = os.path.join(tempfile.gettempdir(), "output.docx")
        create_docx(edited_output, docx_path)
        with open(docx_path, "rb") as f:
            st.download_button("Download as DOCX", f, file_name="output.docx")

        pdf_path = os.path.join(tempfile.gettempdir(), "output.pdf")
        create_pdf(edited_output, pdf_path)
        with open(pdf_path, "rb") as f:
            st.download_button("Download as PDF", f, file_name="output.pdf", mime="application/pdf")
